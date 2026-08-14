from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "docs" / "人事全流程与组织效能管理系统项目规格与实施方案.docx"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("人事全流程与组织效能管理系统项目规格与实施方案")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("需求规格说明书 / 概要设计 / 详细设计 / 测试与验收方案")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(89, 89, 89)

meta = [
    ("项目名称", "人事全流程与组织效能管理系统"),
    ("建议技术栈", "Vue 3 + TypeScript / Spring Boot / PostgreSQL / Redis / MinIO / Docker + Nginx"),
    ("目标读者", "项目负责人、产品经理、前后端开发、测试、运维、业务验收人"),
    ("交付用途", "立项评审、开发启动、迭代执行、测试验收、上线准备"),
]
add_table(doc, ["字段", "内容"], meta, [1.6, 5.8])

heading(doc, "1. 项目概述")
para(doc, "本项目建设一套面向中小型至集团化企业的人事全流程与组织效能管理系统，覆盖组织员工档案、招聘入职、考勤假勤、薪酬绩效、合同附件、流程审批、报表分析、权限审计等核心能力。系统以管理后台为主，兼顾员工自助门户与移动端审批，目标是在统一的数据模型和权限体系下支撑 HR 日常运营、管理决策和合规留痕。")

heading(doc, "2. 建设目标")
bullets(doc, [
    "建立统一员工主数据台账，减少部门、岗位、合同、薪资等信息分散维护的问题。",
    "将入职、转正、调岗、请假、加班、离职等流程线上化，支持可配置审批和全程留痕。",
    "提供考勤、假勤、薪酬、绩效、人员结构等报表，支持导出和权限水印。",
    "通过 RBAC、数据权限、操作审计和附件追溯满足企业内部管控要求。",
    "以容器化方式部署，支持测试、预生产、生产多环境运维和备份恢复。"
])

heading(doc, "3. 用户角色")
roles = [
    ("系统管理员", "维护租户参数、角色权限、字典、审批流程、系统日志。"),
    ("HR 管理员", "维护组织、岗位、员工档案、合同、薪资项、招聘入职与离职。"),
    ("部门主管", "审批请假加班、处理入转调离流程、查看本部门人员与绩效数据。"),
    ("普通员工", "查看个人信息、合同、假勤余额、薪资单，提交请假、补卡等申请。"),
    ("财务专员", "查看薪资汇总、社保公积金、成本中心数据，进行薪酬复核。"),
    ("审计/高管", "查看审计日志、人员结构、组织指标和关键经营报表。"),
]
add_table(doc, ["角色", "职责与权限边界"], roles, [1.5, 5.9])

heading(doc, "4. 功能模块清单")
modules = [
    ("系统管理", "用户、角色、菜单、数据权限、字典、参数、审计日志。", "P0"),
    ("组织与员工档案", "部门、岗位、员工主档、合同、证照、附件、批量导入。", "P0"),
    ("招聘与入职", "招聘需求、候选人、面试、Offer、入职清单。", "P1"),
    ("考勤与假勤", "考勤记录、排班、请假、加班、补卡、假期余额。", "P0"),
    ("薪酬与社保", "薪资项、薪资核算、薪资单、社保公积金、成本中心。", "P1"),
    ("绩效管理", "绩效周期、目标、评分、结果归档。", "P2"),
    ("流程审批", "流程配置、节点条件、待办、审批记录、消息通知。", "P0"),
    ("报表中心", "员工结构、异动、考勤、薪酬、导出、权限水印。", "P1"),
]
add_table(doc, ["模块", "范围", "优先级"], modules, [1.4, 5.1, 0.9])

heading(doc, "5. 页面清单")
pages = [
    ("登录/找回密码", "账号密码、验证码、SSO 入口、登录异常提示。"),
    ("工作台", "待办、快捷入口、员工统计、异常提醒、最近操作。"),
    ("组织架构", "部门树、岗位列表、拖拽排序、启停用。"),
    ("员工档案", "员工列表、详情、合同、证照、附件、变更记录。"),
    ("招聘管理", "招聘需求、候选人、面试、Offer、入职转化。"),
    ("入转调离", "入职清单、转正、调岗、离职交接、流程状态。"),
    ("考勤假勤", "考勤明细、请假、加班、补卡、假期余额。"),
    ("薪酬管理", "薪资项、薪资核算、薪资单发放、社保公积金。"),
    ("绩效管理", "绩效周期、目标填报、评分、结果确认。"),
    ("审批中心", "我的待办、已办、抄送、流程追踪。"),
    ("报表中心", "人员结构、异动、考勤、薪酬、导出任务。"),
    ("系统设置", "用户、角色、菜单、字典、流程、审计日志。"),
]
add_table(doc, ["页面", "主要内容"], pages, [1.8, 5.6])

heading(doc, "6. 数据库表设计")
tables = [
    ("sys_user", "系统用户表", "id, username, password_hash, employee_id, status, last_login_at"),
    ("sys_role", "角色表", "id, code, name, data_scope, status"),
    ("sys_permission", "权限表", "id, type, code, name, path, parent_id"),
    ("org_department", "部门表", "id, parent_id, name, code, manager_id, sort_no, status"),
    ("org_position", "岗位表", "id, dept_id, name, level, headcount, status"),
    ("hr_employee", "员工主档", "id, employee_no, name, gender, dept_id, position_id, hire_date, status"),
    ("hr_employee_contract", "劳动合同", "id, employee_id, contract_no, start_date, end_date, file_id, status"),
    ("recruit_candidate", "候选人", "id, name, phone, source, stage, owner_id"),
    ("attendance_record", "考勤记录", "id, employee_id, work_date, clock_in, clock_out, status"),
    ("leave_request", "请假申请", "id, employee_id, leave_type, start_at, end_at, hours, approval_status"),
    ("payroll_batch", "薪资批次", "id, period, status, generated_by, approved_at"),
    ("payroll_item", "员工薪资明细", "id, batch_id, employee_id, item_code, amount"),
    ("workflow_instance", "流程实例", "id, biz_type, biz_id, starter_id, status, started_at, ended_at"),
    ("workflow_task", "流程任务", "id, instance_id, assignee_id, node_name, action, comment"),
    ("file_object", "文件对象", "id, bucket, object_key, filename, mime_type, size, uploader_id"),
    ("audit_log", "审计日志", "id, actor_id, action, resource_type, resource_id, ip, created_at"),
]
add_table(doc, ["表名", "用途", "关键字段"], tables, [1.8, 1.8, 4.0])

heading(doc, "7. API 接口规划")
apis = [
    ("认证权限", "POST /api/auth/login, GET /api/auth/me, POST /api/auth/logout, GET /api/permissions/tree"),
    ("组织员工", "GET/POST /api/departments, GET/POST /api/positions, GET/POST /api/employees, GET /api/employees/{id}"),
    ("招聘入职", "GET/POST /api/recruit/candidates, POST /api/recruit/offers, POST /api/onboarding/tasks"),
    ("考勤假勤", "GET /api/attendance/records, POST /api/leave-requests, POST /api/overtime-requests"),
    ("薪酬绩效", "POST /api/payroll/batches, GET /api/payroll/slips, POST /api/performance/cycles"),
    ("审批流程", "POST /api/workflows/start, POST /api/workflow-tasks/{id}/approve, GET /api/workflow-tasks/my"),
    ("文件报表", "POST /api/files/upload, GET /api/reports/headcount, GET /api/export-tasks/{id}"),
]
add_table(doc, ["接口域", "核心接口"], apis, [1.7, 5.7])

heading(doc, "8. 前端开发计划")
para(doc, "前端建议采用 Vue 3 + TypeScript + Vite + Pinia + Vue Router + Element Plus。界面方向采用克制的后台管理风格：清晰表格、稳定表单、左侧导航、顶部面包屑、工作台卡片、可响应式折叠的详情抽屉。")
bullets(doc, [
    "第 1 阶段：搭建工程、路由、权限守卫、布局、登录、菜单、请求封装。",
    "第 2 阶段：完成组织员工、档案详情、导入导出、附件上传、审批中心。",
    "第 3 阶段：完成考勤假勤、薪酬绩效、报表中心、移动端审批适配。",
    "第 4 阶段：完成空状态、加载态、错误态、键盘可达性、响应式和性能优化。"
])

heading(doc, "9. 后端开发计划")
para(doc, "后端建议采用 Spring Boot 3 + Spring Security + MyBatis-Plus 或 JPA + PostgreSQL + Redis + MinIO。Spring Boot 在企业人事系统的权限、事务、审计、批处理和团队招聘方面更稳妥。")
bullets(doc, [
    "基础层：统一异常、参数校验、日志、审计、分页、导入导出、文件上传。",
    "权限层：RBAC、菜单权限、按钮权限、数据范围、JWT/SSO 扩展点。",
    "业务层：员工主档、合同、招聘入职、考勤假勤、薪酬绩效、流程审批。",
    "集成层：邮件/企业微信通知、MinIO 文件、Redis 缓存与分布式锁、定时任务。",
    "运维层：健康检查、指标监控、数据库迁移、备份恢复、灰度发布准备。"
])

heading(doc, "10. 测试计划")
tests = [
    ("单元测试", "核心服务、权限判断、流程条件、薪资计算、假期余额计算。"),
    ("接口测试", "认证、员工、考勤、流程、文件、报表接口；覆盖正常和异常路径。"),
    ("前端测试", "表单校验、列表筛选、详情页、审批流、权限可见性、响应式。"),
    ("集成测试", "审批通过后主数据更新、附件上传后回显、薪资生成后导出。"),
    ("性能测试", "10 万员工档案分页查询、批量导入、报表导出、并发审批。"),
    ("安全测试", "越权访问、数据权限绕过、文件下载权限、弱密码、审计完整性。"),
]
add_table(doc, ["测试类型", "重点"], tests, [1.6, 5.8])

heading(doc, "11. 部署计划")
bullets(doc, [
    "环境划分：开发、测试、预生产、生产；各环境数据库、Redis、MinIO 独立。",
    "容器编排：Nginx、前端静态站点、Spring Boot 服务、PostgreSQL、Redis、MinIO 使用 Docker Compose 起步，可平滑迁移 Kubernetes。",
    "发布流程：构建镜像、执行数据库迁移、启动服务、健康检查、烟测、回滚预案。",
    "运维保障：每日数据库备份、MinIO 生命周期策略、日志归档、服务监控、告警通知。"
])

heading(doc, "12. 验收标准")
acceptance = [
    ("功能验收", "P0 功能 100% 通过，P1 功能按里程碑通过，关键流程无阻断缺陷。"),
    ("数据验收", "员工主档、合同、考勤、薪资、审批数据一致，可追溯。"),
    ("权限验收", "角色菜单、按钮、部门数据范围和文件访问权限全部通过。"),
    ("性能验收", "常规列表 2 秒内返回，核心报表 10 秒内生成，批量导入支持 1 万行。"),
    ("安全验收", "无高危漏洞，所有敏感操作有审计日志，薪资和附件访问受控。"),
    ("交付验收", "部署手册、用户手册、测试报告、验收报告、备份恢复方案齐备。"),
]
add_table(doc, ["验收项", "标准"], acceptance, [1.6, 5.8])

heading(doc, "13. 开发里程碑")
milestones = [
    ("M1 立项与原型", "第 1-2 周", "需求确认、原型、技术方案、数据库初版。"),
    ("M2 基础平台", "第 3-4 周", "登录权限、布局、系统管理、基础工程。"),
    ("M3 核心 HR", "第 5-8 周", "组织员工、档案、合同、招聘入职、审批。"),
    ("M4 运营模块", "第 9-12 周", "考勤假勤、薪酬绩效、报表中心、通知。"),
    ("M5 测试上线", "第 13-16 周", "联调、UAT、性能安全、部署上线、验收。"),
]
add_table(doc, ["里程碑", "周期", "交付物"], milestones, [1.7, 1.4, 4.3])

heading(doc, "14. 人员分工建议")
team = [
    ("项目经理", "计划、风险、沟通、验收闭环", "1"),
    ("产品经理", "需求、原型、验收用例、用户培训", "1"),
    ("UI/UX 设计", "后台界面规范、核心页面、响应式方案", "1"),
    ("前端开发", "Vue 工程、页面、组件、权限、联调", "2"),
    ("后端开发", "Spring Boot 服务、数据库、流程、集成", "2"),
    ("测试工程师", "测试计划、用例、自动化、缺陷跟踪", "1"),
    ("DevOps/运维", "容器部署、监控、备份、发布回滚", "0.5-1"),
]
add_table(doc, ["角色", "职责", "建议人数"], team, [1.5, 4.8, 1.1])

heading(doc, "15. 风险与应对措施")
risks = [
    ("需求边界膨胀", "高", "冻结 P0/P1 范围，变更走评审，延期需求进入二期。"),
    ("薪酬规则复杂", "高", "先实现标准薪资项与批次，复杂规则预留公式引擎扩展。"),
    ("历史数据质量差", "中", "提供导入模板、校验报告、数据清洗脚本和灰度导入。"),
    ("权限越权风险", "高", "接口级权限校验、数据范围过滤、审计日志和安全测试。"),
    ("审批流变化频繁", "中", "流程节点和条件配置化，保留版本和实例快照。"),
    ("上线影响业务", "中", "预生产演练、回滚预案、分批导入、上线后双轨观察。"),
]
add_table(doc, ["风险", "等级", "应对措施"], risks, [1.8, 0.9, 4.7])

heading(doc, "附录：配套图示与表格")
bullets(doc, [
    "系统架构图：deliverables/diagram/hr-management-system/system-architecture.svg",
    "业务流程图：deliverables/diagram/hr-management-system/business-flow.svg",
    "数据流图：deliverables/diagram/hr-management-system/data-flow.svg",
    "项目管理工作簿：deliverables/spreadsheets/人事全流程与组织效能管理系统项目管理工作簿.xlsx"
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
